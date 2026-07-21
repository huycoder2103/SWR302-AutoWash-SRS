# -*- coding: utf-8 -*-
"""Sinh cheat sheet thuyết trình 15 phút cho Auto-Wash (SWR302 - SE1916 - Group 2)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x25,0x63,0xEB); GREY=RGBColor(0x55,0x5F,0x70)
doc=Document(); sec=doc.sections[0]
sec.top_margin=sec.bottom_margin=Cm(1.5); sec.left_margin=sec.right_margin=Cm(1.8)
st=doc.styles['Normal']; st.font.name='Arial'; st.font.size=Pt(10)
st.element.rPr.rFonts.set(qn('w:eastAsia'),'Arial')

def shade(cell,color):
    tcPr=cell._tc.get_or_add_tcPr(); el=OxmlElement('w:shd')
    el.set(qn('w:val'),'clear'); el.set(qn('w:fill'),color); tcPr.append(el)

def H(text,size=14,color=NAVY,before=12,after=5):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=color; r.font.name='Arial'; return p

def P(text,size=10,bold=False,italic=False,after=3,indent=0,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after)
    if indent: p.paragraph_format.left_indent=Cm(indent)
    r=p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(size); r.font.name='Arial'
    if color is not None: r.font.color.rgb=color
    return p

def KV(label,text,size=10):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(label+': '); r.bold=True; r.font.size=Pt(size); r.font.name='Arial'; r.font.color.rgb=NAVY
    r2=p.add_run(text); r2.font.size=Pt(size); r2.font.name='Arial'; return p

def B(text,indent=0.6,size=10):
    p=doc.add_paragraph(text,style='List Bullet'); p.paragraph_format.space_after=Pt(2)
    p.paragraph_format.left_indent=Cm(indent)
    for r in p.runs: r.font.size=Pt(size); r.font.name='Arial'
    return p

def TBL(headers,rows,widths,fs=9.5):
    tot=sum(widths); k=17.3/tot; widths=[w*k for w in widths]
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.autofit=False
    hdr=t.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].text=''; p=hdr[i].paragraphs[0]; r=p.add_run(h)
        r.bold=True; r.font.size=Pt(fs); r.font.name='Arial'; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        shade(hdr[i],'1F3864')
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].text=''; p=cs[i].paragraphs[0]; p.paragraph_format.space_after=Pt(1)
            r=p.add_run(str(v)); r.font.size=Pt(fs); r.font.name='Arial'
            if i==0: r.bold=True
    for row in t.rows:
        for i,c in enumerate(row.cells): c.width=Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after=Pt(3)
    return t

def SAY(head, lines):
    """Lời thoại đọc gần như nguyên văn."""
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(head); r.bold=True; r.font.size=Pt(11.5); r.font.color.rgb=NAVY; r.font.name='Arial'
    for ln in lines:
        q=doc.add_paragraph(); q.paragraph_format.left_indent=Cm(0.5); q.paragraph_format.space_after=Pt(4)
        rr=q.add_run(ln); rr.font.size=Pt(10.5); rr.font.name='Arial'

def NOTE(text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5); p.paragraph_format.space_after=Pt(6)
    r=p.add_run('Ghi chú sân khấu: '+text); r.italic=True; r.font.size=Pt(9); r.font.name='Arial'; r.font.color.rgb=GREY

def UCBOX(code,title,owner,minutes):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(4)
    r=p.add_run('%s — %s'%(code,title)); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=NAVY; r.font.name='Arial'
    r2=p.add_run('   [%s · %s]'%(owner,minutes)); r2.bold=True; r2.font.size=Pt(9.5); r2.font.color.rgb=BLUE; r2.font.name='Arial'

def STEPS(rows):
    TBL(['B.','Hệ thống / người dùng làm gì','Vì sao có bước này (câu trả lời nếu bị hỏi)'],rows,[0.9,7.6,8.8])

# ===================== TRANG BÌA =====================
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2)
r=p.add_run('AUTO-WASH — CHEAT SHEET THUYẾT TRÌNH 15 PHÚT'); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=NAVY; r.font.name='Arial'
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2)
r=p.add_run('SRS v2.9 · SWR302 · SE1916 — Group 2 · deck: SWR302-SE1916-Group 2.pptx (14 slide)'); r.font.size=Pt(11); r.font.name='Arial'
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(10)
r=p.add_run('Toàn bộ lời thoại đã viết sẵn, đọc gần như nguyên văn. Nói tiếng Việt, giữ nguyên thuật ngữ tiếng Anh trên slide. Xưng hô: "thầy".')
r.italic=True; r.font.size=Pt(9.5); r.font.name='Arial'; r.font.color.rgb=GREY

H('1. PHÂN CÔNG & ĐỒNG HỒ',size=14,before=4)
TBL(['Slide','Phút','Người','Phần phụ trách'],
[
 ['1','0:00 – 0:15','Nguyễn Hoàng Huy (SE190240)','Mở đầu — chào, giới thiệu nhóm và lộ trình trình bày'],
 ['2','0:15 – 0:45','Nguyễn Hoàng Huy','Background & Purpose — bối cảnh, mục tiêu SRS, chuẩn IEEE 830'],
 ['3','0:45 – 1:05','Nguyễn Hoàng Huy','Technology Architecture'],
 ['4','1:05 – 1:30','Nguyễn Hoàng Huy','Actors & External Systems'],
 ['5','1:30 – 2:10','Nguyễn Hoàng Huy','Context Diagram'],
 ['5','2:10 – 3:45','Nguyễn Hoàng Huy','UC3 Đăng ký tài khoản (đào sâu)'],
 ['6 – 7','3:45 – 4:45','Võ Lê Trung Nguyên (SE190220)','ERD 9 thực thể — bảng và sơ đồ'],
 ['8','4:45 – 5:15','Võ Lê Trung Nguyên','5 phân hệ – 24 Use Case'],
 ['8','5:15 – 6:35','Võ Lê Trung Nguyên','UC8 Đăng ký xe (đào sâu)'],
 ['9','6:35 – 7:20','Nguyễn Lê Thuận (SE190305)','State Transition Diagram'],
 ['10','7:20 – 8:00','Nguyễn Lê Thuận','Data Flow Diagram mức 1'],
 ['10','8:00 – 9:50','Nguyễn Lê Thuận','UC12 Tạo đặt lịch (đào sâu)'],
 ['11','9:50 – 10:15','Nguyễn Thành Đạt (SE190239)','Non-Functional Requirements'],
 ['12 – 13','10:15 – 10:55','Nguyễn Thành Đạt','Dialog Map — khách hàng và nhân viên'],
 ['13','10:55 – 12:55','Nguyễn Thành Đạt','UC17 → UC19 → UC21 Check-in đến Check-out (đào sâu)'],
 ['14','12:55 – 13:15','Nguyễn Thành Đạt','Chốt & mời câu hỏi'],
 ['—','13:15 – 15:00','Cả nhóm','Q&A — ai sở hữu use case nào trả lời use case đó'],
],[1.4,2.3,4.2,9.4])
P('Trình bày đúng thứ tự slide 1 → 14, không nhảy qua nhảy lại. Mỗi người giữ một khối slide liền nhau: Huy 1–5, Nguyên 6–8, Thuận 9–10, Đạt 11–14.',size=9.5,italic=True,after=4)

doc.add_page_break()
H('2. LỜI THOẠI ĐẦY ĐỦ THEO TỪNG SLIDE',size=15,before=0)
P('Phần in thường là lời nói. Ghi chú sân khấu in nghiêng là hướng dẫn thao tác, không đọc.',size=9.5,italic=True,after=6)

SAY('SLIDE 1 — Auto-Wash (Huy · 15 giây)', [
 'Em chào thầy và các bạn. Nhóm em là Group 2, lớp SE1916. Nhóm gồm bốn thành viên: Nguyễn Hoàng Huy, Võ Lê Trung Nguyên, Nguyễn Lê Thuận và Nguyễn Thành Đạt.',
 'Hôm nay nhóm em trình bày tài liệu đặc tả yêu cầu phần mềm — Software Requirements Specification, phiên bản 2.9 — cho hệ thống Auto-Wash, hệ thống quản lý trung tâm rửa xe tự động.',
 'Nhóm em sẽ trình bày theo hành trình của một khách hàng: đăng ký tài khoản, đăng ký xe, đặt lịch, rồi tới cửa hàng check-in và check-out.',
])
NOTE('đứng thẳng, nhìn xuống thầy khi nói câu đầu, chỉ nhìn slide khi bắt đầu slide 2.')

SAY('SLIDE 2 — Background & Purpose (Huy · 30 giây)', [
 'Về bối cảnh. Hiện nay nhiều trung tâm rửa xe vẫn nhận khách và xếp lượt thủ công: khách gọi điện hoặc tới trực tiếp, nhân viên ghi ra giấy. Cách làm này dễ trùng giờ, khách phải chờ lâu, và cửa hàng không biết trước lượng xe trong ngày.',
 'Auto-Wash số hóa toàn bộ quy trình đó, từ lúc khách đặt lịch trên điện thoại cho tới lúc quản lý hàng chờ tại cửa hàng.',
 'Mục tiêu của tài liệu SRS này là chốt phạm vi nghiệp vụ, thống nhất yêu cầu chức năng và phi chức năng giữa các bên, và làm cơ sở cho kiểm thử chấp nhận — User Acceptance Testing.',
 'Sản phẩm là một ứng dụng web: khách hàng dùng Mobile Web, nhân viên dùng màn hình desktop. Tài liệu được viết theo chuẩn IEEE 830.',
])

SAY('SLIDE 3 — Technology Architecture (Huy · 20 giây)', [
 'Về kiến trúc công nghệ. Frontend là React 18 với Vite, dạng single-page application, tối ưu cho điện thoại. Backend là ASP.NET Core 8 cùng Entity Framework Core. Cơ sở dữ liệu là PostgreSQL trên Supabase, bắt buộc kết nối SSL. Hai phía giao tiếp qua REST API dạng JSON.',
 'Có một điểm nhóm em muốn nhấn: phiên đăng nhập được lưu ở phía server. Nhờ vậy backend không giữ trạng thái riêng trong từng instance, và hệ thống có thể chạy nhiều instance phía sau bộ cân bằng tải.',
])
NOTE('slide này đọc nhanh, không dừng lại ở từng dòng.')

SAY('SLIDE 4 — Actors & External Systems (Huy · 25 giây)', [
 'Hệ thống có hai tác nhân người. Khách hàng thì đăng ký tài khoản, quản lý xe của mình, tạo và theo dõi booking trên Mobile Web, và nhận email thông báo. Nhân viên thì xác nhận hoặc hủy booking, check-in, quét biển số, quản lý hàng chờ, check-out, đồng thời quản lý tài khoản và danh mục dịch vụ.',
 'Ngoài ra còn một tác nhân hệ thống: hệ thống tự động gửi email OTP, email xác nhận đặt lịch, email hoàn tất, và tự nhận kết quả nhận dạng biển số.',
 'Ba hệ thống ngoài là Email SMTP qua Gmail và MailKit, Camera LPR, và LPR API của bên thứ ba.',
 'Một điểm quan trọng: khách hàng và nhân viên không dùng chung màn hình nào cả — điều này sẽ thấy rõ ở Dialog Map phần sau.',
])

SAY('SLIDE 5 — Context Diagram (Huy · 40 giây)', [
 'Đây là sơ đồ ngữ cảnh. Ô ở giữa là hệ thống Auto-Wash; mọi thứ nằm ngoài ô này đều nằm ngoài phạm vi nhóm em xây dựng.',
 'Phía Khách hàng: gửi vào hệ thống các yêu cầu đăng ký, đăng nhập, quản lý xe và đặt lịch; nhận về thông tin phiên đăng nhập, danh sách xe, danh sách booking kèm trạng thái.',
 'Phía Nhân viên: gửi vào các thao tác xác nhận và hủy booking, check-in, chuyển công đoạn rửa, check-out; nhận về danh sách hàng chờ trong ngày.',
 'Với Email SMTP, hệ thống gửi ra ba loại email: mã OTP, xác nhận đặt lịch, và thông báo hoàn tất rửa xe.',
 'Với thiết bị nhận dạng, hệ thống gửi yêu cầu chụp ảnh tới Camera LPR; ảnh được chuyển sang LPR API của bên thứ ba và trả về chuỗi biển số.',
 'Sơ đồ này chốt ranh giới hệ thống: hai tác nhân người và ba hệ thống ngoài.',
])
NOTE('chỉ tay theo đúng thứ tự trên: giữa → trái → phải → dưới. Không kể lung tung.')

SAY('SLIDE 5 → UC3 ĐĂNG KÝ TÀI KHOẢN (Huy · 95 giây)', [
 'Use case đầu tiên nhóm em đi sâu là UC3 — Đăng ký tài khoản, vì đây là nơi mọi dữ liệu của hệ thống bắt đầu.',
 'Khách mở màn hình Đăng ký và nhập đầy đủ họ tên, email, số điện thoại và mật khẩu trong một biểu mẫu, rồi gửi đi.',
 'Hệ thống kiểm tra lần lượt ba lớp. Lớp thứ nhất là trường bắt buộc và định dạng. Lớp thứ hai là email đã có ai đăng ký chưa. Lớp thứ ba là số điện thoại đã có ai dùng chưa. Cả email lẫn số điện thoại đều phải duy nhất, vì cả hai đều là định danh đăng nhập hợp lệ.',
 'Chỉ khi ba lớp này đều qua, Notification Service mới sinh mã OTP sáu chữ số, hiệu lực năm phút, và gửi qua Email SMTP. Nhóm em đặt OTP ở sau cùng là có chủ ý: hệ thống chỉ tốn một email khi dữ liệu chắc chắn hợp lệ, chứ không gửi mã cho biểu mẫu sẽ bị từ chối.',
 'Mã OTP được lưu dưới dạng băm theo đúng nguyên tắc áp dụng cho mật khẩu: nếu cơ sở dữ liệu bị lộ thì kẻ tấn công vẫn không đọc được mã đang còn hiệu lực.',
 'Khách nhập mã OTP, màn hình có sẵn nút gửi lại. Hệ thống xác thực và rẽ ba nhánh: nhập sai thì cho nhập lại, mã hết hạn thì quay về bước sinh mã mới, hợp lệ thì đi tiếp.',
 'Sau khi OTP hợp lệ, hệ thống tạo bản ghi Account với mật khẩu đã băm bằng Bcrypt và bản ghi Customer liên kết, trong cùng một giao dịch cơ sở dữ liệu — vì một Account mà không có Customer thì khách không đặt lịch được và cũng không có chỗ lưu thống kê. Cuối cùng khách được đưa về màn hình Đăng nhập.',
 'Ở đây nhóm em tách Account và Customer thành hai bảng: Account giữ phần định danh và bảo mật, dùng chung cho cả nhân viên lẫn khách hàng; còn Customer giữ phần hồ sơ nghiệp vụ mà chỉ khách hàng mới có, như tổng số lượt rửa và tổng chi tiêu.',
])

SAY('CHUYỂN GIAO — Huy → Nguyên', [
 '"Khách đã có tài khoản. Nhưng muốn đặt lịch thì trước hết phải có xe. Em mời bạn Nguyên trình bày mô hình dữ liệu, danh sách chức năng và use case đăng ký xe."',
])

SAY('SLIDE 6 & 7 — ERD 9 thực thể (Nguyên · 60 giây)', [
 'Em xin trình bày phần mô hình dữ liệu. Hệ thống có chín thực thể, và em xin đi theo trục chính thay vì đọc từng bảng.',
 'Một Account tương ứng một hồ sơ Customer. Một Customer có nhiều Vehicle. Một Vehicle có nhiều Booking. Và mỗi Booking khi khách tới cửa hàng sẽ sinh ra tối đa một bản ghi Queue — quan hệ ở đây là không hoặc một, bởi vì có những lượt rửa của khách vãng lai không hề có booking trước.',
 'Có hai chỗ trong thiết kế em muốn nói kỹ hơn.',
 'Thứ nhất là BookingService. Đây là bảng trung gian cho quan hệ nhiều-nhiều giữa Booking và Service, và nó lưu thêm một cột PriceSnapshot, tức giá của dịch vụ tại đúng thời điểm khách đặt. Nếu sau này nhân viên đổi bảng giá thì các hóa đơn cũ vẫn giữ nguyên số tiền đã cam kết với khách.',
 'Thứ hai là OtpVerification. Bảng này cố ý không có khóa ngoại tới Account, vì trong luồng đăng ký, mã OTP được phát hành khi tài khoản còn chưa tồn tại — không có gì để tham chiếu tới. Nó là dữ liệu tạm, định danh bằng email, và bị đánh dấu đã dùng ngay sau khi xác thực.',
 'Ba thực thể còn lại là Service — danh mục dịch vụ có phân loại, giá và thời lượng; Queue — hàng chờ tại cửa hàng; và Notification — thông báo trong ứng dụng.',
])

SAY('SLIDE 8 — 5 phân hệ, 24 Use Case (Nguyên · 30 giây)', [
 'Hệ thống có hai mươi bốn use case, chia theo năm phân hệ: phân hệ chung gồm đăng nhập và đăng xuất; quản lý tài khoản; quản lý phương tiện; dịch vụ và đặt lịch; hàng chờ; và thông báo.',
 'Trong mười lăm phút, nhóm em xin đi sâu bốn use case tạo thành một hành trình hoàn chỉnh: UC3 đăng ký tài khoản, UC8 đăng ký xe, UC12 tạo đặt lịch, và UC17 tới UC21 là từ check-in đến check-out.',
 'Các use case còn lại đều đã được đặc tả đầy đủ luồng chính, luồng thay thế, luồng ngoại lệ và quy tắc nghiệp vụ trong tài liệu SRS ở mục 3, nhóm em xin phép không trình bày hết ạ.',
])
NOTE('nói xong câu cuối thì chuyển slide ngay, không dừng đọc bảng.')

SAY('SLIDE 8 → UC8 ĐĂNG KÝ XE (Nguyên · 80 giây)', [
 'Trong danh sách này, use case em đi sâu là UC8 — Đăng ký xe, vì đây là nơi sinh ra dữ liệu Vehicle, gốc của cả đặt lịch lẫn check-in sau này.',
 'Khách nhập biển số và loại xe. Loại xe để tùy chọn, vì nó không ảnh hưởng tới giá — giá nằm ở bảng Service chứ không nằm ở Vehicle.',
 'Hệ thống kiểm tra định dạng biển số theo biểu thức chính quy cho biển số Việt Nam, và kiểm tra biển số này chưa tồn tại trong tài khoản của khách.',
 'Tiếp theo hệ thống sinh mã OTP sáu chữ số, băm và gửi tới email đã đăng ký. Ở đây nhóm em giữ OTP là có lý do nghiệp vụ: chiếc xe chính là thứ dùng để check-in và nhận dịch vụ tại cửa hàng, nên phải chắc chắn chính chủ tài khoản đang thao tác, chứ không phải người mượn được máy.',
 'Khách nhập OTP và bấm Save. Hệ thống xác thực và rẽ đúng ba nhánh như ở đăng ký: sai mã thì nhập lại, hết hạn thì phát mã mới, hợp lệ thì lưu bản ghi Vehicle với biển số đã được chuẩn hóa.',
 'Bước chuẩn hóa biển số là quy tắc quan trọng nhất của use case này. Khách có thể nhập "51H-123.45", "51H cách 12345", hoặc gõ thường "51h12345". Trong khi đó Camera LPR ở phần sau lại trả về chuỗi thô liền nhau. Nếu hệ thống lưu nguyên văn thì lúc check-in sẽ không tìm ra xe, nhân viên phải nhập tay, và mất luôn liên kết tới booking khách đã đặt.',
 'Vì vậy nhóm em chuẩn hóa ngay tại điểm ghi là use case này, và chuẩn hóa lại tại điểm đọc là check-in và quét biển số, bằng cùng một hàm — để hai đầu chắc chắn khớp nhau.',
])

SAY('CHUYỂN GIAO — Nguyên → Thuận', [
 '"Đã có tài khoản và đã có xe. Tiếp theo là các trạng thái, luồng dữ liệu và use case đặt lịch — use case phức tạp nhất của hệ thống. Em mời bạn Thuận."',
])

SAY('SLIDE 9 — State Transition Diagram (Thuận · 45 giây)', [
 'Em xin bắt đầu bằng sơ đồ chuyển trạng thái, vì nó là luật cho toàn bộ phần đặt lịch và vận hành phía sau. Hệ thống có ba vòng đời cần quản lý.',
 'Thứ nhất là Booking: từ Chờ xác nhận sang Đã xác nhận khi nhân viên duyệt, sang Đã check-in khi xe tới cửa hàng, và sang Hoàn tất khi check-out. Nhánh Đã hủy chỉ đi ra được từ hai trạng thái đầu — nghĩa là booking đã check-in thì không hủy được nữa.',
 'Thứ hai là Queue: từ Đang chờ đi qua các công đoạn rửa rồi tới Hoàn tất, hoặc rẽ sang Đã hủy. Khi rời trạng thái Đang chờ, hệ thống ghi mốc StartedAt; khi vào Hoàn tất thì ghi CompletedAt. Hai mốc này cho ra thời gian phục vụ thực tế.',
 'Thứ ba là OTP: từ Đã phát hành sang Đã dùng, hoặc sang Hết hạn sau năm phút. Không có đường quay lại.',
 'Sơ đồ này chính là luật cho use case chuyển công đoạn: nhân viên chỉ bấm một nút Chuyển tiếp, còn hệ thống tự quyết trạng thái kế tiếp, nên không thể nhảy cóc công đoạn.',
])

SAY('SLIDE 10 — Data Flow Diagram mức 1 (Thuận · 40 giây)', [
 'Đây là sơ đồ luồng dữ liệu mức một. Hệ thống được tách thành năm tiến trình, từ P1 đến P5, tương ứng năm service: Account, Booking, Queue, Notification và LPR. Dữ liệu nằm ở bảy kho, từ D1 đến D7. Hai thực thể ngoài là Khách hàng và Nhân viên.',
 'Em xin đi đúng một luồng đầu-cuối để thầy tiện theo dõi, đó là luồng đặt lịch.',
 'Khách hàng gửi yêu cầu tạo booking vào tiến trình P2 Booking Service. P2 đọc kho D5 Service để lấy giá và thời lượng, sau đó ghi vào kho D3 Booking và kho D4 BookingService.',
 'Ghi xong, P2 gửi một yêu cầu sang P4 Notification Service để gửi email xác nhận. P4 ghi thông báo vào kho D6 Notification, đồng thời gọi ra hệ thống ngoài Email SMTP.',
 'Luồng này khớp đúng với các bước cuối của use case tạo đặt lịch mà em trình bày ngay sau đây.',
])

SAY('SLIDE 10 → UC12 TẠO ĐẶT LỊCH (Thuận · 110 giây)', [
 'UC12 — Tạo đặt lịch là use case phức tạp nhất, nên em xin chia thành ba chặng: chọn, kiểm tra, và ghi.',
 'Chặng thứ nhất là chọn. Hệ thống kiểm tra khách đã có ít nhất một xe hay chưa; nếu chưa thì đưa thẳng sang màn hình thêm xe. Sau đó khách chọn một xe, một dịch vụ chính, các dịch vụ bổ sung nếu muốn, rồi chọn ngày và giờ.',
 'Chặng thứ hai là kiểm tra. Trên sơ đồ hoạt động, năm quy tắc này nằm gọn trong một bước Validate và một điểm quyết định "All rules pass?", vi phạm quy tắc nào thì báo đúng quy tắc đó rồi cho khách sửa. Em xin nói rõ năm quy tắc và rủi ro mà mỗi quy tắc chặn.',
 'Một: chiếc xe phải thuộc về đúng khách đang đăng nhập. Đây là kiểm tra phân quyền ở mức bản ghi, chống trường hợp người dùng sửa mã xe trên request để đặt lịch bằng xe của người khác.',
 'Hai: giờ hẹn phải cách hiện tại ít nhất mười lăm phút và không quá bảy ngày. Mười lăm phút để cửa hàng kịp chuẩn bị; bảy ngày để hàng chờ còn dự báo được và bảng giá không lệch quá xa thời điểm phục vụ.',
 'Ba: khung giờ đó phải đang có ít hơn ba lượt đặt. Đây là ràng buộc năng lực vật lý, vì cửa hàng chỉ phục vụ song song được ba xe. Không có điều kiện này thì hàng chờ sẽ vỡ.',
 'Bốn: không được có booking trùng cho cùng một xe trong cùng khung giờ, chống việc khách bấm hai lần hoặc mở hai tab.',
 'Năm: tất cả dịch vụ được chọn phải đang hoạt động, phòng trường hợp nhân viên vừa vô hiệu hóa một dịch vụ trong lúc khách đang chọn.',
 'Chặng thứ ba là ghi. Hệ thống tính giá cuối bằng tổng giá các dịch vụ đã chọn. Việc tính giá thực hiện ở phía server chứ không lấy số tiền do client gửi lên, vì nếu tin client thì khách có thể sửa được giá.',
 'Hệ thống hiển thị màn hình tóm tắt gồm xe, dịch vụ, ngày giờ và thành tiền để khách xem lại. Màn hình này là yêu cầu về tính khả dụng trong phần phi chức năng, và cũng là điểm cuối cùng khách còn quay lại được.',
 'Khách bấm xác nhận. Hệ thống lưu bản ghi Booking và các dòng BookingService trong cùng một giao dịch cơ sở dữ liệu duy nhất — vì một booking không có dịch vụ nào là dữ liệu vô nghĩa; nếu lỗi giữa chừng thì rollback toàn bộ.',
 'Cuối cùng, Notification Service tạo thông báo trong ứng dụng rồi gửi email xác nhận. Hai bước này nằm sau giao dịch: nếu SMTP lỗi thì booking vẫn hợp lệ và khách vẫn xem được trên ứng dụng, hệ thống chỉ cần thử gửi lại email.',
])

SAY('CHUYỂN GIAO — Thuận → Đạt', [
 '"Booking đã được ghi nhận và email xác nhận đã gửi đi. Phần yêu cầu chất lượng và toàn bộ vận hành tại cửa hàng, em mời bạn Đạt."',
])

SAY('SLIDE 11 — Non-Functional Requirements (Đạt · 25 giây)', [
 'Về yêu cầu phi chức năng, nhóm em đặt các ngưỡng đo được chứ không nói chung chung.',
 'Hiệu năng: API phản hồi dưới hai giây ở phân vị chín mươi lăm, tạo booking dưới ba giây, email OTP tới trong vòng sáu mươi giây, kết quả nhận dạng biển số trong năm giây, và chịu được năm mươi phiên đồng thời.',
 'Bảo mật: mật khẩu băm bằng Bcrypt, OTP băm và dùng một lần, phiên đăng nhập quản lý ở phía server, thông tin bí mật để trong biến môi trường, và bắt buộc HTTPS.',
 'Khả dụng: thời gian hoạt động tối thiểu chín mươi chín phần trăm trong khung bảy giờ sáng tới mười giờ tối, bảo trì chỉ làm ngoài giờ, và có phương án nhập tay thay thế khi thiết bị nhận dạng biển số gặp sự cố.',
])

SAY('SLIDE 12 & 13 — Dialog Map (Đạt · 40 giây)', [
 'Đây là hai luồng màn hình của hệ thống.',
 'Bên khách hàng: Đăng nhập, vào Trang chủ, xem Dịch vụ, sang màn hình Tạo đặt lịch, qua màn hình Tóm tắt, xác nhận, rồi về màn hình Đặt lịch của tôi để theo dõi trạng thái. Màn hình Tóm tắt là bắt buộc, đúng như phần tạo đặt lịch bạn Thuận vừa trình bày.',
 'Bên nhân viên: Đăng nhập là vào thẳng màn hình Quản lý hàng chờ. Từ đó nhân viên mở Check-in, trong đó có nút quét biển số; rồi chuyển công đoạn; rồi Check-out; và luôn quay về màn hình hàng chờ. Hàng chờ là màn hình trung tâm, nên nhân viên không phải nhớ đường đi trong lúc đông khách.',
 'Hai luồng này tách biệt hoàn toàn, đúng với phân quyền trong sơ đồ use case.',
])

SAY('SLIDE 13 → UC17, UC19, UC21 — CHECK-IN ĐẾN CHECK-OUT (Đạt · 120 giây)', [
 'Phần cuối là toàn bộ vận hành tại cửa hàng, gồm ba use case nối tiếp nhau.',
 'Bắt đầu là UC17 — Check-in. Nhân viên bấm nút Quét biển số. Hệ thống gửi yêu cầu chụp ảnh tới Camera LPR, ảnh được chuyển sang LPR API của bên thứ ba, và trả về chuỗi biển số. Kết quả này chỉ được điền sẵn vào biểu mẫu, nhân viên vẫn phải xác nhận — vì nhận dạng sai là chuyện bình thường nên nhóm em không để máy tự quyết.',
 'Biển số nhận được sẽ được chuẩn hóa bằng đúng hàm đã dùng lúc đăng ký xe. Đây chính là lý do bước chuẩn hóa ở UC8 quan trọng: hai đầu ghi và đọc dùng chung một chuẩn thì mới khớp nhau được.',
 'Hệ thống tìm xe theo biển số đã chuẩn hóa. Nếu tìm thấy thì lấy luôn tên khách hàng, khách quen không phải khai lại. Nếu không tìm thấy thì dùng tên do nhân viên nhập — nghĩa là hệ thống vẫn phục vụ được khách vãng lai không có tài khoản, đúng như thực tế cửa hàng.',
 'Tiếp theo hệ thống tìm booking của xe này hẹn trong ngày hôm nay. Nếu có thì liên kết lượt hàng chờ với booking đó và chuyển booking sang trạng thái Đã check-in — từ giây phút này khách xem được vị trí hàng chờ của mình ngay trên điện thoại.',
 'Cuối cùng hệ thống tạo bản ghi Queue ở trạng thái Waiting, xếp theo thứ tự đến trước phục vụ trước chứ không theo giờ hẹn.',
 'Sang UC19 — chuyển công đoạn. Xe đi lần lượt qua bốn công đoạn: LPR, Washing, Add-on và Drying. Nhân viên chỉ bấm chuyển tiếp, hệ thống kiểm tra chuyển trạng thái có đúng thứ tự công đoạn hay không rồi mới cập nhật, nên không thể nhảy cóc.',
 'Cuối cùng là UC21 — Check-out. Hệ thống chuyển cả bản ghi Queue lẫn Booking sang trạng thái Hoàn tất, hai bảng phải đồng bộ, nếu lệch thì màn hình khách và màn hình nhân viên sẽ hiển thị hai chuyện khác nhau.',
 'Ở đây sơ đồ có một nhánh riêng cho khách vãng lai: nếu lượt này không gắn với booking nào thì nhân viên phải nhập số tiền cho các dịch vụ đã thực hiện, vì không có giá chốt sẵn. Còn khách có booking thì hệ thống lấy luôn giá đã chốt lúc đặt — đây là chỗ cột PriceSnapshot bạn Nguyên nói lúc nãy phát huy tác dụng.',
 'Sau đó hệ thống cập nhật hồ sơ khách hàng: tăng tổng số lượt lên một, cộng thêm số tiền vào tổng chi tiêu, và ghi lại lần ghé gần nhất. Đây là nơi duy nhất trong toàn hệ thống cập nhật thống kê khách hàng, và chỉ cập nhật khi xe thực sự đã rửa xong, chứ không cộng từ lúc đặt lịch.',
 'Cuối cùng, hệ thống kiểm tra khách có tài khoản hay không. Có tài khoản thì gửi email hoàn tất; khách vãng lai không có email nên luồng kết thúc ngay tại đó.',
])

SAY('SLIDE 14 — Thank You (Đạt · 20 giây)', [
 'Tóm lại, tài liệu SRS phiên bản 2.9 của nhóm em đặc tả đầy đủ hai mươi bốn use case trên năm phân hệ, kèm chín sơ đồ thiết kế đã được rà soát đồng bộ với nhau, tổng cộng khoảng sáu mươi lăm trang theo chuẩn IEEE 830.',
 'Phần trình bày của nhóm em xin được kết thúc tại đây. Nhóm em cảm ơn thầy và các bạn đã lắng nghe, và xin sẵn sàng nhận câu hỏi ạ.',
])
NOTE('cả bốn người đứng lên phía trước, không ai ngồi xuống trong lúc Q&A.')

doc.add_page_break()
H('3. TRA CỨU NHANH KHI BỊ HỎI — 4 USE CASE ĐÀO SÂU',size=15,before=0)
P('Phần này không đọc trên sân khấu. Đây là bảng tra khi thầy hỏi ngược lại một bước cụ thể.',size=9.5,italic=True,after=6)

UCBOX('UC3','Đăng ký tài khoản','Huy','tra cứu')
KV('Tác nhân','Khách hàng (Mobile Web)')
KV('Tiền điều kiện','Chưa có tài khoản; có email truy cập được')
KV('Hậu điều kiện','Có Account (mật khẩu băm) + Customer liên kết; OTP đã đánh dấu đã dùng')
P('')
STEPS([
 ['1','Mở màn hình Đăng ký, nhập họ tên, email, số điện thoại, mật khẩu trong MỘT biểu mẫu.','Sơ đồ gốc gộp một bước — khách điền một lần, không tách email ra trước.'],
 ['2','Kiểm tra trường bắt buộc và định dạng (Input valid?).','Chặn tại client và server; sai thì báo ngay tại trường, chưa gọi cơ sở dữ liệu.'],
 ['3','Tra email trong hệ thống (Email registered?).','Email là định danh đăng nhập nên phải duy nhất toàn hệ thống.'],
 ['4','Kiểm tra số điện thoại (Phone in use?).','SĐT cũng là định danh đăng nhập hợp lệ ở UC1.'],
 ['5','Sinh OTP 6 số, hiệu lực 5 phút, lưu dạng băm.','Đặt SAU ba lớp kiểm tra: không tốn email cho dữ liệu chắc chắn bị từ chối.'],
 ['6','Gửi email OTP qua SMTP (mẫu ET1).','—'],
 ['7','Khách nhập OTP (có nút gửi lại) → Verify OTP.','Ba nhánh: wrong → nhập lại · expired → quay về bước sinh mã · valid → đi tiếp.'],
 ['8','Tạo Account (Bcrypt) + Customer trong MỘT giao dịch.','Account không có Customer thì không đặt lịch được, không có chỗ lưu thống kê.'],
 ['9','Chuyển sang màn hình Đăng nhập.','—'],
])
TBL(['Thầy hỏi','Trả lời'],
[
 ['Vì sao OtpVerification không có khóa ngoại?','Vì ở chính use case này, OTP được phát hành khi tài khoản chưa tồn tại nên không có gì để tham chiếu. Nó là dữ liệu tạm, định danh bằng email, và bị đánh dấu đã dùng sau khi xác thực ạ.'],
 ['Đăng ký cần OTP mà đăng nhập lại không cần, có mâu thuẫn không?','Dạ không ạ. OTP ở đăng ký để chứng minh email có thật và thuộc về người này — đây là lần duy nhất hệ thống chưa biết họ là ai. Khi đăng nhập thì đã có mật khẩu băm và session, thêm OTP mỗi lần chỉ tăng ma sát mà không tăng an toàn tương xứng. Nhóm em giữ OTP đúng ở ba chỗ nhạy cảm: đăng ký, đổi mật khẩu và đăng ký xe.'],
 ['Vì sao tách Account và Customer?','Account giữ định danh và bảo mật, dùng chung cho cả nhân viên lẫn khách hàng. Customer giữ hồ sơ nghiệp vụ chỉ khách hàng mới có. Nếu gộp thì tài khoản nhân viên phải mang những cột thống kê luôn rỗng ạ.'],
 ['Email gửi chậm hoặc lỗi thì sao?','Yêu cầu phi chức năng quy định email OTP tới trong 60 giây; khách có thể yêu cầu gửi lại. Mã cũ vẫn hết hạn đúng 5 phút của nó, không bị kéo dài ạ.'],
],[4.6,13.0])

doc.add_page_break()
UCBOX('UC8','Đăng ký xe','Nguyên','tra cứu')
KV('Tác nhân','Khách hàng'); KV('Tiền điều kiện','Đã đăng nhập; email đã xác minh')
KV('Hậu điều kiện','Có Vehicle gắn CustomerId, biển số đã chuẩn hóa')
P('')
STEPS([
 ['1','Nhập biển số + loại xe (tùy chọn).','Loại xe tùy chọn vì không ảnh hưởng giá — giá nằm ở Service.'],
 ['2','Kiểm tra định dạng biển số (regex biển số Việt Nam).','Chặn rác đầu vào; dữ liệu bẩn khiến LPR không bao giờ khớp.'],
 ['3','Biển số chưa tồn tại trong tài khoản này.','Chống trùng trong phạm vi một khách hàng.'],
 ['4','Sinh OTP 6 số (băm, 5 phút) gửi email đã đăng ký — SAU khi biển số đã hợp lệ và không trùng.','Xe là thứ dùng để check-in và nhận dịch vụ, phải chắc chắn chính chủ thao tác; và không gửi mã cho biển số sai định dạng.'],
 ['5','Verify OTP — ba nhánh: wrong → nhập lại · expired → phát mã mới · valid → đi tiếp.','Tái sử dụng nguyên cơ chế OTP của UC3.'],
 ['6','Lưu Vehicle với biển số CHUẨN HÓA (CBR4): viết hoa, bỏ chấm/gạch/khoảng trắng.','Điểm ghi và điểm đọc dùng cùng một hàm thì mới khớp — xem UC17, UC18.'],
 ['7','Hiển thị MSG17.','Mọi thông báo có mã MSG để đối chiếu khi làm UAT.'],
])
TBL(['Thầy hỏi','Trả lời'],
[
 ['Hai khách khác nhau đăng ký cùng một biển số?','Ràng buộc duy nhất chỉ trong phạm vi một tài khoản, vì thực tế xe có thể đổi chủ hoặc dùng chung trong gia đình. Khi check-in, nếu biển số khớp nhiều xe thì hệ thống ưu tiên bản ghi có booking hôm nay ạ.'],
 ['Xóa xe đã có booking được không?','Dạ không — UC9 chặn bằng MSG6. Booking lịch sử phải tra ngược được về xe; nếu cho xóa thì dữ liệu doanh thu và thống kê sẽ mồ côi ạ.'],
 ['Vì sao đăng ký xe lại cần OTP?','Vì đây là thao tác gắn một tài sản vật lý vào tài khoản. Người mượn được phiên đăng nhập vẫn không thêm được xe nếu không đọc được email của chủ tài khoản ạ.'],
],[4.6,13.0])

doc.add_page_break()
UCBOX('UC12','Tạo đặt lịch','Thuận','tra cứu')
KV('Tác nhân','Khách hàng'); KV('Tiền điều kiện','Đã đăng nhập và có ít nhất 1 xe')
KV('Hậu điều kiện','Booking (Chờ xác nhận) + các dòng BookingService; đã gửi email xác nhận')
P('')
P('5 điều kiện kiểm tra — mỗi điều kiện chặn một rủi ro khác nhau',bold=True,after=2)
TBL(['#','Điều kiện','Chặn rủi ro gì'],
[
 ['1','Xe thuộc về khách đang đăng nhập (CBR5).','Chống sửa mã xe trên request để đặt bằng xe người khác — phân quyền mức bản ghi.'],
 ['2','Giờ hẹn ≥ 15 phút và ≤ 7 ngày.','15 phút để cửa hàng kịp chuẩn bị; 7 ngày để hàng chờ dự báo được và giá không lệch xa.'],
 ['3','Khung giờ đang có < 3 lượt (CBR7).','Ràng buộc năng lực vật lý — cửa hàng chỉ phục vụ song song 3 xe.'],
 ['4','Không trùng booking cùng xe cùng khung giờ.','Chống double-booking do bấm 2 lần hoặc mở 2 tab.'],
 ['5','Mọi dịch vụ đang IsActive.','Chống đặt dịch vụ vừa bị vô hiệu ở UC11 giữa lúc khách đang chọn.'],
],[0.8,6.2,10.6])
P('Chặng ghi',bold=True,after=2)
TBL(['B.','Hành động','Vì sao'],
[
 ['10','FinalPrice = tổng giá dịch vụ đã chọn (CBR6).','Tính ở SERVER, không lấy số tiền client gửi lên — tin client thì khách sửa được giá.'],
 ['11–12','Màn hình Tóm tắt → khách bấm Xác nhận.','Yêu cầu usability; điểm cuối cùng khách còn quay lại được.'],
 ['13','Lưu Booking + BookingService trong MỘT giao dịch (CBR9).','Booking không có dịch vụ là dữ liệu vô nghĩa; lỗi giữa chừng thì rollback toàn bộ.'],
 ['14','Tạo thông báo + gửi email xác nhận (ET2).','Gửi email NGOÀI giao dịch — SMTP lỗi thì booking vẫn hợp lệ, chỉ thử gửi lại.'],
 ['15','Hiển thị MSG14.','—'],
],[1.0,6.0,10.6])
TBL(['Thầy hỏi','Trả lời'],
[
 ['Hai khách cùng đặt slot cuối cùng đúng một lúc?','Kiểm tra CBR7 và việc ghi Booking nằm trong cùng một giao dịch nên người thứ hai sẽ thất bại ở bước kiểm tra và được mời chọn khung giờ khác. Nếu tải cao hơn thì cần khóa mức hàng hoặc ràng buộc duy nhất trên cặp khung giờ và vị trí ạ.'],
 ['Vì sao Booking tạo ra ở trạng thái Chờ xác nhận?','Vì cửa hàng vẫn cần một người xác nhận năng lực thực tế ở UC14. Việc chuyển trạng thái tuân đúng sơ đồ chuyển trạng thái, không nhảy cóc ạ.'],
 ['Email xác nhận lỗi thì booking có bị hủy không?','Dạ không. Email nằm ngoài giao dịch, booking đã hợp lệ trong cơ sở dữ liệu và khách vẫn xem được ở UC13. Thông báo trong ứng dụng là kênh chính, email là kênh phụ ạ.'],
 ['Vì sao giới hạn 7 ngày?','Là quyết định nghiệp vụ trong phạm vi đồ án, để hàng chờ dự báo được và giá không lệch quá xa thời điểm phục vụ. Con số này được ghi thành quy tắc nên đổi được mà không ảnh hưởng thiết kế ạ.'],
 ['PriceSnapshot để làm gì?','Giữ giá tại thời điểm đặt. Nếu sau này nhân viên tăng giá ở UC11 thì hóa đơn cũ vẫn giữ nguyên số tiền đã cam kết; nếu chỉ tham chiếu Service.Price thì hóa đơn quá khứ sẽ tự đổi theo bảng giá ạ.'],
],[4.6,13.0])

doc.add_page_break()
UCBOX('UC17 → UC19 → UC21','Check-in · Chuyển công đoạn · Check-out','Đạt','tra cứu')
KV('Tác nhân','Nhân viên · hệ thống ngoài Camera LPR + 3rd-party LPR API')
KV('Hậu điều kiện','Queue và Booking đều Hoàn tất; thống kê khách đã cập nhật; đã gửi email hoàn tất')
P('')
STEPS([
 ['1','Bấm Quét biển số → Camera LPR chụp → 3rd-party LPR API trả chuỗi biển số (UC18).','LPR chỉ ĐIỀN SẴN, người vẫn xác nhận — nhận dạng sai là bình thường nên không cho máy tự quyết.'],
 ['2','Chuẩn hóa biển số (CBR4) — cùng hàm với UC8.','Đây là lý do UC8 phải chuẩn hóa: hai đầu ghi và đọc dùng chung chuẩn thì mới khớp.'],
 ['3','Tìm Vehicle theo biển số đã chuẩn hóa.','—'],
 ['4a','Tìm thấy → lấy tên khách hàng liên kết.','Khách quen không phải khai lại thông tin.'],
 ['4b','Không thấy → dùng tên nhân viên nhập (khách vãng lai).','Hệ thống không bắt buộc có tài khoản mới được rửa xe — đúng thực tế cửa hàng.'],
 ['5','Tìm Booking của xe hẹn hôm nay (Chờ xác nhận / Đã xác nhận).','Nối lượt tại chỗ với lịch đặt online.'],
 ['6','Có booking → liên kết Queue và đặt Booking = Đã check-in.','Từ đây khách xem được vị trí hàng chờ trên điện thoại (UC13).'],
 ['7','Tạo Queue: Đang chờ, Position cuối hàng.','Thứ tự phục vụ FIFO theo CheckInAt (CBR8), không theo giờ hẹn.'],
 ['UC19','Advance stages: LPR → Washing → Add-on → Drying; hệ thống kiểm tra chuyển trạng thái đúng thứ tự công đoạn.','Sai thứ tự thì bị từ chối — chống nhảy cóc và giữ dữ liệu thời gian sạch.'],
 ['UC21','Check-out: khách vãng lai thì nhân viên nhập tiền, khách có booking thì lấy giá đã chốt → Set Completed → cập nhật TotalVisits, TotalSpend, LastVisitAt → gửi email hoàn tất (ET3) NẾU khách có tài khoản.','Đây là NƠI DUY NHẤT cập nhật thống kê khách hàng, và chỉ khi xe thực sự rửa xong.'],
])
TBL(['Thầy hỏi','Trả lời'],
[
 ['Camera LPR đọc sai biển số thì hệ thống hỏng à?','Dạ không. LPR chỉ điền sẵn, nhân viên sửa và xác nhận trước khi check-in. Yêu cầu phi chức năng cũng đã ghi rõ có phương án nhập tay khi LPR lỗi hoặc chậm quá 5 giây ạ.'],
 ['Khách không đặt trước có phục vụ được không?','Dạ được — bước 4b, hệ thống tạo lượt hàng chờ với tên do nhân viên nhập và để trống mã booking. Vì vậy quan hệ Booking–Queue trên ERD là không hoặc một, chứ không phải một-một ạ.'],
 ['Vì sao xếp theo giờ check-in mà không theo giờ hẹn?','Vì thứ tự phục vụ phải phản ánh ai đang thật sự có mặt. Khách đặt 9 giờ nhưng 10 giờ mới tới thì không thể chen trước người đã đến từ 9 giờ rưỡi. Giờ hẹn chỉ dùng để giữ năng lực theo quy tắc CBR7 ạ.'],
 ['Hủy lượt giữa chừng thì dữ liệu ra sao?','UC22: kiểm tra lượt chưa Hoàn tất, chuyển Queue sang Đã hủy, và nếu có mã booking thì đồng bộ Booking cũng Đã hủy. Thống kê khách không được cộng vì chỉ UC21 mới cộng ạ.'],
],[4.6,13.0])

doc.add_page_break()
H('4. SỐ LIỆU PHẢI THUỘC (cả nhóm)',before=0)
TBL(['Nhóm','Con số / từ khóa'],
[
 ['Tài liệu','SRS v2.9 · ~65 trang · IEEE 830 · 24 UC · 5 phân hệ · 9 sơ đồ · 23 mockup · 32 mã MSG · 3 mẫu email ET'],
 ['Kiến trúc','React 18 + Vite · ASP.NET Core 8 + EF Core · PostgreSQL (Supabase, SSL) · REST JSON · session phía server'],
 ['Dữ liệu','9 thực thể: Account, Customer, Vehicle, Booking, BookingService, Service, Queue, Notification, OtpVerification'],
 ['OTP','6 chữ số · băm · 5 phút · dùng một lần · CHỈ ở UC3 đăng ký, UC5 đổi mật khẩu, UC8 đăng ký xe'],
 ['Đặt lịch','≥ 15 phút và ≤ 7 ngày · tối đa 3 lượt/khung giờ (CBR7) · FinalPrice = tổng giá (CBR6) · ghi trong 1 giao dịch (CBR9)'],
 ['Quy tắc chung','CBR4 chuẩn hóa biển số · CBR5 kiểm tra quyền sở hữu bản ghi · CBR8 hàng chờ FIFO theo CheckInAt'],
 ['NFR','API < 2s (P95) · tạo booking < 3s · email OTP < 60s · LPR < 5s · 50 phiên đồng thời · uptime ≥ 99% (07:00–22:00)'],
 ['Ngoài phạm vi','Tích điểm · voucher · campaign · admin dashboard · Google OAuth · đa chi nhánh'],
],[3.2,14.4])

H('5. BỊ HỎI UC KHÔNG TRÌNH BÀY — TRẢ LỜI 2 CÂU')
TBL(['UC','Người','Trả lời'],
[
 ['UC1 Đăng nhập','Huy','Nhập email hoặc số điện thoại cùng mật khẩu, hệ thống kiểm tra tài khoản còn hoạt động, đối chiếu chuỗi băm, tạo session rồi điều hướng theo vai trò. Đăng nhập không dùng OTP, lý do em đã trình bày ở phần đăng ký ạ.'],
 ['UC2 Đăng xuất','Huy','Hệ thống xóa session ở phía máy chủ rồi xóa cookie xác thực, sau đó đưa về màn hình đăng nhập. Xóa ở server chứ không chỉ xóa phía client ạ.'],
 ['UC5 Đổi mật khẩu','Huy','Gửi OTP tới email đã đăng ký, xác thực OTP rồi băm và cập nhật mật khẩu mới. Giữ OTP ở đây vì đây là thao tác có thể chiếm quyền tài khoản ạ.'],
 ['UC4 Cập nhật hồ sơ','Nguyên','Sửa họ tên và số điện thoại, hệ thống kiểm tra tên không rỗng và số điện thoại mới phải duy nhất toàn hệ thống rồi mới lưu ạ.'],
 ['UC6 Quản lý tài khoản','Nguyên','Nhân viên tìm kiếm theo tên, email hoặc số điện thoại rồi bật tắt trạng thái hoạt động. Nhóm em vô hiệu hóa chứ không xóa, để giữ toàn vẹn booking lịch sử ạ.'],
 ['UC7 / UC9 Xem, xóa xe','Nguyên','Xem thì lọc theo khách đang đăng nhập. Xóa thì kiểm tra quyền sở hữu và kiểm tra xe đã có booking chưa; đã có thì chặn bằng MSG6 ạ.'],
 ['UC10 / UC11 Dịch vụ','Thuận','Khách chỉ thấy dịch vụ đang hoạt động, phân nhóm theo Cơ bản, Cao cấp, Đặc biệt và Add-On. Nhân viên thấy cả dịch vụ đã vô hiệu và có thể tạo, sửa, kiểm tra giá và thời lượng là số dương ạ.'],
 ['UC13 Xem đặt lịch','Thuận','Truy vấn booking theo khách, và nếu booking đang hoạt động thì lấy thêm bản ghi hàng chờ để hiện vị trí và công đoạn hiện tại theo thời gian thực ạ.'],
 ['UC14 / UC15 Xác nhận, hủy','Thuận','Chỉ chuyển được từ đúng trạng thái nguồn theo sơ đồ chuyển trạng thái, kèm thông báo cho khách, mã MSG24 và MSG25 ạ.'],
 ['UC16 Xem hàng chờ','Đạt','Gộp các bản ghi hàng chờ trong ngày với các booking hẹn hôm nay chưa check-in, khử trùng lặp, rồi sắp xếp theo thời điểm check-in tăng dần ạ.'],
 ['UC18 Quét biển số','Đạt','Camera chụp, LPR API bên thứ ba trả chuỗi biển số, hệ thống chuẩn hóa rồi điền vào biểu mẫu check-in để nhân viên xác nhận ạ.'],
 ['UC20 / UC22 Cập nhật, hủy lượt','Đạt','Cập nhật cho phép sửa trạng thái và ghi chú của nhân viên. Hủy lượt thì chuyển hàng chờ sang Đã hủy và đồng bộ booking nếu có liên kết ạ.'],
 ['UC23 / UC24 Thông báo','Đạt','Liệt kê thông báo theo khách, mới nhất trước, ba loại là đặt lịch, hoàn tất và hệ thống. Bấm vào thì đánh dấu đã đọc sau khi kiểm tra quyền sở hữu ạ.'],
],[2.9,1.3,13.4])

H('6. CHECKLIST TRƯỚC KHI VÀO PHÒNG')
for x in [
 'Mở sẵn ba thứ: file SWR302-SE1916-Group 2.pptx, bản SRS PDF (để nhảy đúng trang khi bị hỏi use case ngoài danh sách), và trang index.html chứa 9 sơ đồ tương tác.',
 'Mỗi người đọc thành tiếng phần lời thoại của mình ít nhất hai lần và bấm giờ. 15 phút rất chật, quá giờ là mất phần kết.',
 'Tập riêng ba câu chuyển giao — đây là chỗ nhóm dễ đứt mạch nhất.',
 'Người đang nói cầm chuột, người kế tiếp đứng sẵn bên cạnh.',
 'Ai sở hữu use case nào trả lời use case đó. Bị hỏi ngoài phần mình thì nói "Phần này bạn ... phụ trách ạ" rồi chuyển, không đoán.',
 'Bị hỏi use case không trình bày: dùng bảng mục 5, trả lời hai câu, rồi mở SRS mục 3 nếu thầy muốn xem chi tiết.',
 'Câu chốt cuối do Đạt nói: 24 use case, 9 sơ đồ đồng bộ, khoảng 65 trang theo chuẩn IEEE 830.',
]: B(x)

import sys
doc.save(sys.argv[1] if len(sys.argv)>1 else 'CheatSheet_15phut_G2.docx')
print('saved')
