# -*- coding: utf-8 -*-
"""Luồng hoạt động (Activity Flow) cho 4 use case đào sâu: UC3, UC8, UC12, UC17→UC21."""
import os, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x25,0x63,0xEB); GREY=RGBColor(0x55,0x5F,0x70)
RED=RGBColor(0xB4,0x23,0x18)
doc=Document(); sec=doc.sections[0]
sec.top_margin=sec.bottom_margin=Cm(1.5); sec.left_margin=sec.right_margin=Cm(1.8)
st=doc.styles['Normal']; st.font.name='Arial'; st.font.size=Pt(10)
st.element.rPr.rFonts.set(qn('w:eastAsia'),'Arial')

def shade(cell,color):
    tcPr=cell._tc.get_or_add_tcPr(); el=OxmlElement('w:shd')
    el.set(qn('w:val'),'clear'); el.set(qn('w:fill'),color); tcPr.append(el)

def H(text,size=14,color=NAVY,before=12,after=5):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=color; r.font.name='Arial'; return p

def P(text,size=10,bold=False,italic=False,after=3,indent=0,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after)
    if indent: p.paragraph_format.left_indent=Cm(indent)
    r=p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(size); r.font.name='Arial'
    if color is not None: r.font.color.rgb=color
    return p

def FLOW(lines):
    """Mỗi dòng: (mức thụt, ký hiệu, nội dung, kiểu). kiểu: 'a' action, 'd' decision, 'e' end/exception, 's' start."""
    for lvl,sym,txt,kind in lines:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1)
        p.paragraph_format.left_indent=Cm(0.4+lvl*0.75)
        r=p.add_run(sym+'  '); r.font.name='Consolas'; r.font.size=Pt(9.5)
        r2=p.add_run(txt); r2.font.name='Arial'; r2.font.size=Pt(9.5)
        if kind=='d': r.font.color.rgb=BLUE; r2.bold=True; r2.font.color.rgb=BLUE
        elif kind=='e': r.font.color.rgb=RED; r2.font.color.rgb=RED
        elif kind=='s': r.font.color.rgb=NAVY; r2.bold=True; r2.font.color.rgb=NAVY
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def TBL(headers,rows,widths,fs=9):
    tot=sum(widths); k=17.3/tot; widths=[w*k for w in widths]
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.autofit=False
    hdr=t.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].text=''; p=hdr[i].paragraphs[0]; r=p.add_run(h)
        r.bold=True; r.font.size=Pt(fs); r.font.name='Arial'; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        shade(hdr[i],'1F3864')
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].text=''; p=cs[i].paragraphs[0]; p.paragraph_format.space_after=Pt(1)
            r=p.add_run(str(v)); r.font.size=Pt(fs); r.font.name='Arial'
            if i==0: r.bold=True
    for row in t.rows:
        for i,c in enumerate(row.cells): c.width=Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after=Pt(3)
    return t

def LANES(rows):
    TBL(['#','Khách hàng / Nhân viên','Hệ thống (service)','Hệ thống ngoài / Kho dữ liệu'],rows,[0.8,5.2,6.6,4.7])


from PIL import Image as _PILImage
DGDIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', 'diagrams', 'diagram picture', 'Activity')
def DIAGRAM(fname, caption):
    fp = os.path.join(DGDIR, fname)
    if not os.path.exists(fp):
        P('[thiếu ảnh: %s]' % fname, italic=True); return
    w0, h0 = _PILImage.open(fp).size
    maxw, maxh = 15.5, 22.5
    w = min(maxw, maxh * w0 / h0); h = w * h0 / w0
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(fp, width=Cm(w), height=Cm(h))
    q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = q.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.name = 'Arial'; r.font.color.rgb = GREY
    doc.add_page_break()

# ================= BÌA =================
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2)
r=p.add_run('AUTO-WASH — LUỒNG HOẠT ĐỘNG 4 USE CASE TRÌNH BÀY'); r.bold=True; r.font.size=Pt(17); r.font.color.rgb=NAVY; r.font.name='Arial'
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2)
r=p.add_run('UC3 Đăng ký tài khoản · UC8 Đăng ký xe · UC12 Tạo đặt lịch · UC17→UC19→UC21 Check-in đến Check-out'); r.font.size=Pt(11); r.font.name='Arial'
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(10)
r=p.add_run('SRS v2.9 · SWR302 · SE1916 — Group 2'); r.italic=True; r.font.size=Pt(9.5); r.font.name='Arial'; r.font.color.rgb=GREY

H('CÁCH ĐỌC KÝ HIỆU',size=12,before=2)
TBL(['Ký hiệu','Nghĩa'],
[
 ['●','Điểm bắt đầu / điểm kết thúc của luồng (initial node / final node).'],
 ['□','Hành động (action) — một bước xử lý của người dùng hoặc của hệ thống.'],
 ['◇','Điểm quyết định (decision node) — luồng rẽ nhánh theo điều kiện.'],
 ['├ Đ / └ S','Nhánh Đúng và nhánh Sai đi ra từ điểm quyết định ngay phía trên.'],
 ['↺','Quay lại một bước trước đó (vòng lặp).'],
 ['⊗','Kết thúc theo nhánh lỗi / ngoại lệ (flow final node).'],
],[2.0,15.3])
P('Mỗi use case gồm bốn phần: (1) ảnh Activity Diagram gốc của nhóm, (2) luồng hoạt động viết bằng lời theo ĐÚNG sơ đồ đó, (3) bảng phân làn theo tác nhân, (4) bảng nhánh rẽ và ngoại lệ. Toàn bộ nội dung bám theo diagrams/ActivityDiagrams.drawio — không tự thêm bước nào ngoài sơ đồ.',size=9.5,italic=True,after=4)

doc.add_page_break()

