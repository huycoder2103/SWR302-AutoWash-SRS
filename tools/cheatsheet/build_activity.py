# -*- coding: utf-8 -*-
"""Luồng hoạt động (Activity Flow) cho 4 use case đào sâu: UC3, UC8, UC12, UC17→UC21."""
import os, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x25,0x63,0xEB); GREY=RGBColor(0x55,0x5F,0x70)
RED=RGBColor(0xB4,0x23,0x18)
doc=Document(); sec=doc.sections[0]
sec.top_margin=sec.bottom_margin=Cm(1.5); sec.left_margin=sec.right_margin=Cm(1.8)
st=doc.styles['Normal']; st.font.name='Arial'; st.font.size=Pt(10)
st.element.rPr.rFonts.set(qn('w:eastAsia'),'Arial')

def shade(cell,color):
    tcPr=cell._tc.get_or_add_tcPr(); el=OxmlElement('w:shd')
    el.set(qn('w:val'),'clear'); el.set(qn('w:fill'),color); tcPr.append(el)

def H(text,size=14,color=NAVY,before=12,after=5):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=color; r.font.name='Arial'; return p

def P(text,size=10,bold=False,italic=False,after=3,indent=0,color=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after)
    if indent: p.paragraph_format.left_indent=Cm(indent)
    r=p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(size); r.font.name='Arial'
    if color is not None: r.font.color.rgb=color
    return p

def FLOW(lines):
    """Mỗi dòng: (mức thụt, ký hiệu, nội dung, kiểu). kiểu: 'a' action, 'd' decision, 'e' end/exception, 's' start."""
    for lvl,sym,txt,kind in lines:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1)
        p.paragraph_format.left_indent=Cm(0.4+lvl*0.75)
        r=p.add_run(sym+'  '); r.font.name='Consolas'; r.font.size=Pt(9.5)
        r2=p.add_run(txt); r2.font.name='Arial'; r2.font.size=Pt(9.5)
        if kind=='d': r.font.color.rgb=BLUE; r2.bold=True; r2.font.color.rgb=BLUE
        elif kind=='e': r.font.color.rgb=RED; r2.font.color.rgb=RED
        elif kind=='s': r.font.color.rgb=NAVY; r2.bold=True; r2.font.color.rgb=NAVY
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def TBL(headers,rows,widths,fs=9):
    tot=sum(widths); k=17.3/tot; widths=[w*k for w in widths]
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.autofit=False
    hdr=t.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].text=''; p=hdr[i].paragraphs[0]; r=p.add_run(h)
        r.bold=True; r.font.size=Pt(fs); r.font.name='Arial'; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        shade(hdr[i],'1F3864')
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].text=''; p=cs[i].paragraphs[0]; p.paragraph_format.space_after=Pt(1)
            r=p.add_run(str(v)); r.font.size=Pt(fs); r.font.name='Arial'
            if i==0: r.bold=True
    for row in t.rows:
        for i,c in enumerate(row.cells): c.width=Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after=Pt(3)
    return t

def LANES(rows):
    TBL(['#','Khách hàng / Nhân viên','Hệ thống (service)','Hệ thống ngoài / Kho dữ liệu'],rows,[0.8,5.2,6.6,4.7])


from PIL import Image as _PILImage
DGDIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', 'diagrams', 'diagram picture', 'Activity')
def DIAGRAM(fname, caption):
    fp = os.path.join(DGDIR, fname)
    if not os.path.exists(fp):
        P('[thiếu ảnh: %s]' % fname, italic=True); return
    w0, h0 = _PILImage.open(fp).size
    maxw, maxh = 15.5, 22.5
    w = min(maxw, maxh * w0 / h0); h = w * h0 / w0
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(fp, width=Cm(w), height=Cm(h))
    q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = q.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.name = 'Arial'; r.font.color.rgb = GREY
    doc.add_page_break()

# ================= BÌA =================
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2)
r=p.add_run('AUTO-WASH — LUỒNG HOẠT ĐỘNG 4 USE CASE TRÌNH BÀY'); r.bold=True; r.font.size=Pt(17); r.font.color.rgb=NAVY; r.font.name='Arial'
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2)
r=p.add_run('UC3 Đăng ký tài khoản · UC8 Đăng ký xe · UC12 Tạo đặt lịch · UC17→UC19→UC21 Check-in đến Check-out'); r.font.size=Pt(11); r.font.name='Arial'
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(10)
r=p.add_run('SRS v2.9 · SWR302 · SE1916 — Group 2'); r.italic=True; r.font.size=Pt(9.5); r.font.name='Arial'; r.font.color.rgb=GREY

H('CÁCH ĐỌC KÝ HIỆU',size=12,before=2)
TBL(['Ký hiệu','Nghĩa'],
[
 ['●','Điểm bắt đầu / điểm kết thúc của luồng (initial node / final node).'],
 ['□','Hành động (action) — một bước xử lý của người dùng hoặc của hệ thống.'],
 ['◇','Điểm quyết định (decision node) — luồng rẽ nhánh theo điều kiện.'],
 ['├ Đ / └ S','Nhánh Đúng và nhánh Sai đi ra từ điểm quyết định ngay phía trên.'],
 ['↺','Quay lại một bước trước đó (vòng lặp).'],
 ['⊗','Kết thúc theo nhánh lỗi / ngoại lệ (flow final node).'],
],[2.0,15.3])
P('Mỗi use case gồm bốn phần: (1) ảnh Activity Diagram gốc của nhóm, (2) luồng hoạt động viết bằng lời theo ĐÚNG sơ đồ đó, (3) bảng phân làn theo tác nhân, (4) bảng nhánh rẽ và ngoại lệ. Toàn bộ nội dung bám theo diagrams/ActivityDiagrams.drawio — không tự thêm bước nào ngoài sơ đồ.',size=9.5,italic=True,after=4)

doc.add_page_break()

# ================= UC3 =================
H('1. UC3 — ĐĂNG KÝ TÀI KHOẢN',size=15,before=0)
P('Nguồn: ActivityDiagrams.drawio › trang "1. Account" › nhánh A. Register',size=9.5,italic=True,after=2)
P('Tác nhân: Khách hàng · Hệ thống: Account Service, Notification Service · Hệ thống ngoài: Email SMTP',size=9.5,italic=True,after=6)
DIAGRAM('ActivityDiagrams-1. Account.drawio.png','Hình 1: Activity Diagram — phân hệ Account. UC3 là nhánh A. Register (hàng trên cùng).')
H('1.1 Luồng hoạt động',size=12)
FLOW([
 (0,'●','BẮT ĐẦU','s'),
 (0,'□','Open registration form — khách mở màn hình Đăng ký','a'),
 (0,'□','Enter full name, email, phone, password — khách nhập TẤT CẢ thông tin trong một biểu mẫu','a'),
 (0,'□','Validate required fields and formats — kiểm tra trường bắt buộc và định dạng','a'),
 (0,'◇','Input valid?','d'),
 (1,'└ no','Báo lỗi tại trường tương ứng  ↺ quay lại biểu mẫu đăng ký','e',),
 (0,'□','Check email in system — tra email trong bảng Account','a'),
 (0,'◇','Email registered?','d'),
 (1,'└ yes','Báo email đã được đăng ký  ↺ quay lại biểu mẫu','e'),
 (0,'□','Check phone uniqueness — kiểm tra số điện thoại chưa ai dùng','a'),
 (0,'◇','Phone in use?','d'),
 (1,'└ yes','Báo số điện thoại đã được sử dụng  ↺ quay lại biểu mẫu','e'),
 (0,'□','Generate OTP (valid 5 min) — CHỈ sinh OTP sau khi dữ liệu đã hợp lệ và không trùng','a'),
 (0,'□','Send OTP email — gửi mã qua Email SMTP','a'),
 (0,'□','Enter OTP (resend available) — khách nhập mã, có nút gửi lại','a'),
 (0,'□','Verify OTP','a'),
 (0,'◇','OTP result?','d'),
 (1,'├ wrong','Sai mã  ↺ quay lại bước nhập OTP','e'),
 (1,'├ expired','Hết hạn  ↺ quay lại bước Generate OTP để phát mã mới','e'),
 (1,'└ valid','đi tiếp','a'),
 (0,'□','Create Account + Customer record — tạo hai bản ghi trong một giao dịch','a'),
 (0,'□','Continue to Login screen','a'),
 (0,'●','KẾT THÚC','s'),
])
P('Điểm thiết kế đáng nói khi thuyết trình: OTP nằm SAU toàn bộ khâu kiểm tra dữ liệu. Hệ thống chỉ tốn một email khi biểu mẫu đã hợp lệ, email chưa ai đăng ký và số điện thoại chưa ai dùng — không gửi mã cho dữ liệu chắc chắn sẽ bị từ chối.',size=9.5,italic=True,after=6)
H('1.2 Phân làn theo tác nhân',size=12)
LANES([
 ['1','Mở màn hình Đăng ký, nhập họ tên / email / SĐT / mật khẩu','—','—'],
 ['2','Gửi biểu mẫu','Account Service kiểm tra trường bắt buộc và định dạng','—'],
 ['3','—','Tra email và số điện thoại đã tồn tại chưa','Đọc Account'],
 ['4','—','Notification Service sinh OTP (hiệu lực 5 phút)','Ghi OtpVerification · gửi Email SMTP'],
 ['5','Nhận email, nhập OTP (có thể bấm gửi lại)','Verify OTP','Đọc OtpVerification'],
 ['6','—','Tạo Account + Customer trong một giao dịch','Ghi Account · ghi Customer'],
 ['7','Được đưa về màn hình Đăng nhập','—','—'],
])
H('1.3 Nhánh rẽ và ngoại lệ',size=12)
TBL(['Điểm quyết định','Nhánh','Hệ thống xử lý'],
[
 ['Input valid?','no','Báo lỗi ngay tại trường sai, giữ nguyên dữ liệu đã nhập, chưa gọi tới cơ sở dữ liệu.'],
 ['Email registered?','yes','Dừng, báo email đã tồn tại. Email là định danh đăng nhập nên phải duy nhất.'],
 ['Phone in use?','yes','Dừng, báo số điện thoại đã được dùng — số điện thoại cũng là định danh đăng nhập hợp lệ.'],
 ['OTP result?','wrong','Cho nhập lại mã, không phát mã mới.'],
 ['OTP result?','expired','Quay về bước sinh OTP để phát mã mới (nút resend).'],
 ['OTP result?','valid','Tạo Account và Customer trong cùng một giao dịch.'],
],[3.4,2.2,11.7])

doc.add_page_break()

# ================= UC8 =================
H('2. UC8 — ĐĂNG KÝ XE',size=15,before=0)
P('Nguồn: ActivityDiagrams.drawio › trang "2. Vehicle" › nhánh B. Register vehicle (OTP)',size=9.5,italic=True,after=2)
P('Tác nhân: Khách hàng · Hệ thống: Account Service, Notification Service · Hệ thống ngoài: Email SMTP',size=9.5,italic=True,after=6)
DIAGRAM('ActivityDiagrams-2. Vehicle.drawio.png','Hình 2: Activity Diagram — phân hệ Vehicle. UC8 là nhánh B. Register vehicle (OTP), nhánh giữa.')
H('2.1 Luồng hoạt động',size=12)
FLOW([
 (0,'●','BẮT ĐẦU — khách đã đăng nhập, mở màn hình Thêm xe','s'),
 (0,'□','Enter plate and vehicle type — nhập biển số và loại xe','a'),
 (0,'□','Validate VN plate format — kiểm tra định dạng biển số Việt Nam','a'),
 (0,'◇','Format valid?','d'),
 (1,'└ no','Báo sai định dạng  ↺ quay lại bước nhập biển số','e'),
 (0,'□','Check plate not already registered — kiểm tra biển số chưa đăng ký','a'),
 (0,'◇','Duplicate?','d'),
 (1,'└ yes','Báo trùng biển số  ↺ quay lại bước nhập biển số','e'),
 (0,'□','Generate OTP — sinh mã sau khi biển số đã hợp lệ và không trùng','a'),
 (0,'□','Send OTP email — gửi tới email đã đăng ký của khách','a'),
 (0,'□','Enter OTP, press Save','a'),
 (0,'□','Verify OTP','a'),
 (0,'◇','OTP result?','d'),
 (1,'├ wrong','Sai mã  ↺ quay lại bước nhập OTP','e'),
 (1,'├ expired','Hết hạn  ↺ quay lại bước Generate OTP để phát mã mới','e'),
 (1,'└ valid','đi tiếp','a'),
 (0,'□','Save vehicle (normalized plate) — lưu Vehicle với biển số đã chuẩn hóa (CBR4)','a'),
 (0,'●','KẾT THÚC','s'),
])
P('Chuẩn hóa biển số nằm ngay ở bước lưu: viết hoa, bỏ dấu chấm, dấu gạch và khoảng trắng. Đây là điều kiện để bước "Normalize plate, look up vehicle" khi check-in (trang Queue) tìm ra đúng chiếc xe này.',size=9.5,italic=True,after=6)
H('2.2 Phân làn theo tác nhân',size=12)
LANES([
 ['1','Nhập biển số và loại xe','Kiểm tra định dạng biển số Việt Nam','—'],
 ['2','—','Kiểm tra biển số chưa đăng ký','Đọc Vehicle'],
 ['3','—','Notification Service sinh OTP','Ghi OtpVerification · gửi Email SMTP'],
 ['4','Nhập OTP, bấm Save','Verify OTP','Đọc OtpVerification'],
 ['5','Thấy xe trong danh sách Xe của tôi','Lưu Vehicle với biển số đã chuẩn hóa','Ghi Vehicle'],
])
H('2.3 Nhánh rẽ và ngoại lệ',size=12)
TBL(['Điểm quyết định','Nhánh','Hệ thống xử lý'],
[
 ['Format valid?','no','Báo sai định dạng và cho sửa — chưa gửi OTP, không tốn email.'],
 ['Duplicate?','yes','Báo biển số đã có trong tài khoản, quay lại bước nhập.'],
 ['OTP result?','wrong','Cho nhập lại mã.'],
 ['OTP result?','expired','Quay về bước sinh OTP để phát mã mới.'],
 ['OTP result?','valid','Lưu Vehicle với biển số đã chuẩn hóa.'],
],[3.4,2.2,11.7])
P('Hai nhánh còn lại trên cùng trang sơ đồ (dùng khi thầy hỏi thêm): A. View vehicles — nếu khách chưa có xe nào thì hiện màn hình rỗng kèm gợi ý thêm xe. C. Delete vehicle — kiểm tra quyền sở hữu, nếu không phải chủ xe thì từ chối; nếu xe đã có lịch sử booking thì chặn xóa để giữ lịch sử.',size=9.5,italic=True,after=4)

doc.add_page_break()

# ================= UC12 =================
H('3. UC12 — TẠO ĐẶT LỊCH',size=15,before=0)
P('Nguồn: ActivityDiagrams.drawio › trang "3. Booking" › nhánh A. Create booking (Customer)',size=9.5,italic=True,after=2)
P('Tác nhân: Khách hàng · Hệ thống: Booking Service, Notification Service · Hệ thống ngoài: Email SMTP',size=9.5,italic=True,after=6)
DIAGRAM('ActivityDiagrams-3. Booking.drawio.png','Hình 3: Activity Diagram — phân hệ Booking. UC12 là nhánh A. Create booking (Customer), hàng trên cùng.')
H('3.1 Luồng hoạt động',size=12)
FLOW([
 (0,'●','BẮT ĐẦU','s'),
 (0,'□','Open booking screen — khách mở màn hình đặt lịch','a'),
 (0,'□','Check customer has vehicles','a'),
 (0,'◇','Any vehicle?','d'),
 (1,'└ no','Prompt to register a vehicle — điều hướng sang UC8  ⊗ kết thúc luồng đặt lịch','e'),
 (0,'□','Select vehicle, services, date and time — chọn xe, dịch vụ và ngày giờ','a'),
 (0,'□','Validate: ownership, time window, slot capacity, no duplicate, active services','a'),
 (1,'','Năm quy tắc được kiểm tra trong CÙNG một bước: quyền sở hữu xe (CBR5) · khung thời gian hợp lệ · khung giờ còn chỗ (CBR7) · không trùng booking · dịch vụ còn hoạt động','a'),
 (0,'◇','All rules pass?','d'),
 (1,'└ no','Hiển thị đúng quy tắc bị vi phạm  ↺ quay lại bước chọn xe / dịch vụ / ngày giờ','e'),
 (0,'□','Calculate price (sum of services) — tính tổng giá các dịch vụ đã chọn (CBR6)','a'),
 (0,'□','Review summary, press Confirm — khách xem màn hình tóm tắt','a'),
 (0,'◇','Confirm?','d'),
 (1,'└ no - adjust','Khách quay lại chỉnh sửa  ↺ về bước chọn xe / dịch vụ / ngày giờ','e'),
 (0,'□','Save booking in one transaction (Pending) — ghi Booking và BookingService trong một giao dịch','a'),
 (0,'□','Create confirmation notification — tạo thông báo trong ứng dụng','a'),
 (0,'□','Send booking confirmation email — gửi email xác nhận qua SMTP','a'),
 (0,'●','KẾT THÚC — booking ở trạng thái Pending, chờ nhân viên xác nhận (UC14)','s'),
])
P('Lưu ý khi trình bày: trên sơ đồ, năm quy tắc nghiệp vụ nằm gọn trong một hành động "Validate" và một điểm quyết định "All rules pass?" chứ không tách thành năm nhánh riêng. Nếu thầy hỏi cụ thể từng quy tắc thì đọc bảng 3.3.',size=9.5,italic=True,after=6)
H('3.2 Phân làn theo tác nhân',size=12)
LANES([
 ['1','Mở màn hình đặt lịch','Booking Service kiểm tra khách đã có xe chưa','Đọc Vehicle'],
 ['2','Chọn xe, dịch vụ, ngày giờ','—','Đọc Service'],
 ['3','Gửi yêu cầu','Kiểm tra 5 quy tắc trong một bước','Đọc Vehicle · Booking · Service'],
 ['4','—','Tính tổng giá dịch vụ (CBR6)','—'],
 ['5','Xem tóm tắt, bấm Confirm','—','—'],
 ['6','—','Ghi Booking (Pending) + BookingService trong một giao dịch','Ghi Booking · ghi BookingService'],
 ['7','Nhận thông báo và email xác nhận','Notification Service tạo thông báo và gửi email','Ghi Notification · gọi Email SMTP'],
])
H('3.3 Năm quy tắc trong bước Validate',size=12)
TBL(['Quy tắc trên sơ đồ','Mã','Chặn rủi ro gì'],
[
 ['ownership','CBR5','Xe phải thuộc về khách đang đăng nhập — chống sửa mã xe trên request để đặt bằng xe người khác.'],
 ['time window','—','Giờ hẹn phải cách hiện tại ít nhất 15 phút và không quá 7 ngày.'],
 ['slot capacity','CBR7','Mỗi khung giờ tối đa 3 lượt — ràng buộc năng lực phục vụ của cửa hàng.'],
 ['no duplicate','—','Không có booking trùng cho cùng xe trong cùng khung giờ — chặn double-booking.'],
 ['active services','—','Dịch vụ đã chọn phải còn hoạt động, phòng khi nhân viên vừa vô hiệu hóa ở UC11.'],
],[3.6,1.4,12.3])
H('3.4 Nhánh rẽ và ngoại lệ',size=12)
TBL(['Điểm quyết định','Nhánh','Hệ thống xử lý'],
[
 ['Any vehicle?','no','Điều hướng sang màn hình thêm xe (UC8) thay vì báo lỗi cụt.'],
 ['All rules pass?','no','Hiển thị đúng quy tắc bị vi phạm, giữ nguyên các lựa chọn khác để khách chỉ sửa phần sai.'],
 ['Confirm?','no - adjust','Quay lại bước chọn — chưa ghi gì xuống cơ sở dữ liệu.'],
 ['Save booking','—','Booking và BookingService ghi trong MỘT giao dịch (CBR9); lỗi thì rollback, không để lại booking rỗng.'],
],[3.4,2.2,11.7])
P('Ba nhánh còn lại trên cùng trang (dùng khi bị hỏi): B. View bookings — nếu có booking đang hoạt động và đã check-in thì hiện vị trí hàng chờ và công đoạn rửa theo thời gian thực. C. Confirm booking (Staff) — kiểm tra lại trạng thái vẫn là Pending rồi mới chuyển Confirmed. D. Cancel booking (Staff) — chỉ hủy được khi trạng thái là Pending hoặc Confirmed, có hộp thoại xác nhận trước khi hủy.',size=9.5,italic=True,after=4)

doc.add_page_break()

# ================= UC17 -> UC21 =================
H('4. UC17 → UC19 → UC21 — CHECK-IN ĐẾN CHECK-OUT',size=15,before=0)
P('Nguồn: ActivityDiagrams.drawio › trang "4. Queue" › nhánh A. Check-in to check-out',size=9.5,italic=True,after=2)
P('Tác nhân: Nhân viên · Hệ thống: LPR Service, Queue Service, Booking Service, Notification Service · Hệ thống ngoài: Camera LPR, 3rd-party LPR API, Email SMTP',size=9.5,italic=True,after=6)
DIAGRAM('ActivityDiagrams-4. Queue.drawio.png','Hình 4: Activity Diagram — phân hệ Queue. UC17, UC19 và UC21 nằm trên nhánh A. Check-in to check-out (hàng trên cùng).')
H('4.1 Luồng hoạt động',size=12)
FLOW([
 (0,'●','BẮT ĐẦU — Vehicle arrives at center','s'),
 (0,'□','Press Scan plate (LPR) — nhân viên bấm quét biển số','a'),
 (0,'□','Camera captures; LPR API extracts text — camera chụp, LPR API bên thứ ba trích xuất chuỗi biển số','a'),
 (0,'◇','Recognized in 5 s?','d'),
 (1,'└ no','Enter plate manually — nhân viên nhập tay rồi nhập lại luồng chính','a'),
 (0,'□','Normalize plate, look up vehicle + booking — chuẩn hóa biển số rồi tra xe và booking','a'),
 (0,'◇','Booking found?','d'),
 (1,'└ yes','Link booking, CheckedIn, pre-fill — liên kết booking, đặt trạng thái Đã check-in, điền sẵn dịch vụ','a'),
 (0,'□','Create queue entry (Waiting), FIFO — tạo lượt hàng chờ, xếp theo thứ tự đến trước phục vụ trước','a'),
 (0,'□','Advance stages: LPR, Washing, Add-on, Drying — UC19, nhân viên chuyển lần lượt qua các công đoạn','a'),
 (0,'□','Press Check-out — UC21','a'),
 (0,'◇','Walk-in without booking?','d'),
 (1,'└ yes','Enter amount for services — nhân viên nhập số tiền cho khách vãng lai rồi nhập lại luồng chính','a'),
 (0,'□','Set Completed, update visit stats — đóng lượt và cập nhật TotalVisits, TotalSpend, LastVisitAt','a'),
 (0,'◇','Customer has account?','d'),
 (1,'└ no','Không gửi email — khách vãng lai không có tài khoản, không có địa chỉ email  ⊗','e'),
 (0,'□','Send wash-complete email — gửi email hoàn tất qua SMTP','a'),
 (0,'●','KẾT THÚC','s'),
])
P('Hai chi tiết trên sơ đồ dễ bị bỏ sót khi thuyết trình: khách vãng lai không có booking thì nhân viên phải NHẬP SỐ TIỀN thủ công lúc check-out, và khách không có tài khoản thì hệ thống bỏ qua bước gửi email hoàn tất.',size=9.5,italic=True,after=6)
H('4.2 Phân làn theo tác nhân',size=12)
LANES([
 ['1','Bấm Scan plate','LPR Service gửi yêu cầu chụp','Camera LPR → 3rd-party LPR API'],
 ['2','Nhập tay nếu quá 5 giây','Chuẩn hóa biển số, tra Vehicle và Booking','Đọc Vehicle · đọc Booking'],
 ['3','—','Liên kết booking, đặt CheckedIn, điền sẵn dịch vụ','Cập nhật Booking'],
 ['4','—','Queue Service tạo lượt Waiting, xếp FIFO','Ghi Queue'],
 ['5','Bấm chuyển từng công đoạn: LPR → Washing → Add-on → Drying','Cập nhật Queue.Status theo đúng thứ tự công đoạn','Cập nhật Queue'],
 ['6','Bấm Check-out; nhập tiền nếu là khách vãng lai','Đóng lượt, cập nhật thống kê khách hàng','Cập nhật Queue · Booking · Customer'],
 ['7','—','Gửi email hoàn tất nếu khách có tài khoản','Ghi Notification · gọi Email SMTP'],
])
H('4.3 Nhánh rẽ và ngoại lệ',size=12)
TBL(['Điểm quyết định','Nhánh','Hệ thống xử lý'],
[
 ['Recognized in 5 s?','no','Nhân viên nhập biển số thủ công — LPR chỉ hỗ trợ, không phải điều kiện bắt buộc.'],
 ['Booking found?','yes','Liên kết lượt hàng chờ với booking, chuyển booking sang Đã check-in, điền sẵn dịch vụ đã đặt.'],
 ['Booking found?','no','Tạo lượt hàng chờ độc lập — khách vãng lai. Đây là lý do quan hệ Booking–Queue là 0..1.'],
 ['Walk-in without booking?','yes','Nhân viên nhập số tiền cho các dịch vụ đã thực hiện, vì không có FinalPrice sẵn từ booking.'],
 ['Walk-in without booking?','no','Lấy giá đã chốt trong booking để đóng lượt.'],
 ['Customer has account?','no','Bỏ qua bước gửi email — kết thúc luồng tại đây.'],
],[3.4,2.2,11.7])
P('Hai nhánh còn lại trên cùng trang: B. Cancel queue entry (UC22) — nếu lượt đã Completed thì từ chối, ngược lại chuyển lượt sang Cancelled và nếu có booking liên kết thì đặt booking Cancelled theo. C. Update queue entry (UC20) — nếu chưa có bản ghi thì tạo mới; nếu có thì kiểm tra "Transition hợp lệ?" theo đúng thứ tự công đoạn, sai thứ tự thì từ chối, đúng thì cập nhật trạng thái và ghi chú của nhân viên.',size=9.5,italic=True,after=4)

H('5. BỐN LUỒNG NÀY NỐI VỚI NHAU NHƯ THẾ NÀO',size=14)
TBL(['Từ luồng','Sinh ra dữ liệu','Được dùng ở luồng sau'],
[
 ['UC3 Đăng ký','Account + Customer (một giao dịch)','UC8 cần CustomerId để gắn xe; UC21 cập nhật TotalVisits và TotalSpend lên chính Customer này.'],
 ['UC8 Đăng ký xe','Vehicle với biển số đã chuẩn hóa','UC12 kiểm tra quyền sở hữu xe; bước "Normalize plate, look up vehicle" khi check-in tra đúng chuỗi biển số này.'],
 ['UC12 Đặt lịch','Booking (Pending) + BookingService','UC14 xác nhận; check-in tra "Booking found?" để liên kết và điền sẵn dịch vụ; check-out lấy giá đã chốt.'],
 ['UC17→21 Vận hành','Queue + cập nhật Booking, Customer, Notification','Khép vòng: khách xem vị trí hàng chờ ở UC13 và nhận thông báo ở UC23.'],
],[3.0,4.4,9.9])

doc.save(sys.argv[1] if len(sys.argv)>1 else 'LuongHoatDong_4UC.docx')
print('saved')
